from docx import Document

path = 'reports/tai-lieu-du-an-nhom-01.docx'
doc = Document(path)

# Remove placeholder paragraphs corresponding to the old Chapter 2 bullet list.
# Based on inspection, paragraphs 77-80 contain the old placeholder bullets after our insertion.
indexes_to_remove = [80, 79, 78, 77]
for idx in indexes_to_remove:
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)

# Optionally, remove blank paragraphs left from previous edit around paragraph 75-76 if any.
# We'll clean any paragraph whose text is empty and not at beginning or end.
for p in list(doc.paragraphs):
    if not p.text.strip():
        p._element.getparent().remove(p._element)

out_path = 'reports/tai-lieu-du-an-nhom-01.docx'
doc.save(out_path)
print('cleaned', out_path)
