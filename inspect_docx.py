from docx import Document
p='reports/tai-lieu-du-an-nhom-01.docx'
doc=Document(p)
print('paragraphs', len(doc.paragraphs))
for i,p in enumerate(doc.paragraphs):
    text=p.text.strip()
    if text:
        print(i, text[:220])
