from docx import Document

path = 'reports/tai-lieu-du-an-nhom-01.docx'
doc = Document(path)

# Paragraph indices from inspection: 71..75 correspond to Chapter 2 placeholder content
start_idx = 71
end_idx = 75

new_text_blocks = [
    'Chương 2. Xây dựng pipeline và tiền xử lý dữ liệu',
    'Đây là chương trọng tâm thể hiện năng lực xây dựng một pipeline dữ liệu hoàn chỉnh, từ việc tích hợp nhiều bảng có quan hệ đến việc làm sạch và kiểm soát chất lượng dữ liệu trước khi đưa vào phân tích và huấn luyện mô hình.',
    '2.1. Quy trình tích hợp dữ liệu: Dữ liệu được thu thập từ nhiều bảng khác nhau, trong đó bảng application_train là bảng trung tâm và các bảng phụ như bureau, previous_application, installments_payments, POS_CASH_balance, credit_card_balance cung cấp thông tin lịch sử tín dụng của khách hàng. Vì số lượng dòng lớn và dữ liệu phân tán trên nhiều bảng, nhóm đã thiết kế pipeline theo hướng phân tầng, trong đó các bảng phụ được tổng hợp trước ở mức khách hàng rồi mới nối vào bảng chính bằng phép LEFT JOIN để tạo thành bảng application_flat. Cách làm này giúp giảm tải cho môi trường tính toán, tối ưu bộ nhớ và giữ được tính nhất quán của dữ liệu.',
    '2.2. Làm sạch dữ liệu chuyên sâu: Sau khi tích hợp dữ liệu, nhóm thực hiện xử lý các vấn đề về khuyết thiếu, giá trị sai logic và outlier. Không phải giá trị thiếu nào cũng là lỗi; với các biến từ bảng phụ, nhiều giá trị NULL phản ánh đúng thực tế rằng khách hàng chưa từng có lịch sử tương ứng. Vì vậy, nhóm không áp dụng cách điền giá trị trung bình hoặc trung vị một cách máy móc mà phân nhóm và xử lý theo ngữ cảnh nghiệp vụ. Với các biến tiền tệ và các đặc trưng lịch sử, nhóm cũng áp dụng các phép biến đổi như cắt theo phân vị hoặc log-transform để giảm ảnh hưởng của outlier mà vẫn giữ được ý nghĩa thực tế của dữ liệu.',
    '2.3. Kiểm thử dữ liệu và kiểm soát chất lượng: Một phần quan trọng của pipeline là kiểm tra chất lượng dữ liệu sau mỗi bước biến đổi. Nhóm đã xây dựng các kịch bản kiểm thử để xác minh số dòng, số cột, tính duy nhất của khóa khách hàng, kiểu dữ liệu và việc loại bỏ các giá trị sai logic. Việc này giúp phát hiện sớm các lỗi trong join, aggregate hoặc chuyển đổi dữ liệu, tránh để sai lệch bị kéo dài vào các bước phân tích và mô hình hóa sau này. Đây là yếu tố then chốt để đảm bảo dữ liệu đầu vào có chất lượng cao và có thể tin cậy cho việc dự báo rủi ro tín dụng.'
]

# Remove the existing paragraphs in the range, then insert new ones.
# Keep the first paragraph as heading and body paragraphs after it.
for _ in range(end_idx - start_idx + 1):
    p = doc.paragraphs[start_idx]
    # Remove from document by clearing its text, then continue; this keeps indices stable enough for loop.
    p.text = ''
    p.clear()

# Rebuild the content by inserting paragraphs at the original start index.
# We will insert from the end backwards to preserve positions.
for block in reversed(new_text_blocks):
    doc.paragraphs[start_idx].insert_paragraph_before(block)

# Fix the first paragraph style to heading-like and rest to normal.
for i in range(start_idx, start_idx + len(new_text_blocks)):
    para = doc.paragraphs[i]
    if i == start_idx:
        para.style = doc.styles['Heading 1']
    else:
        para.style = doc.styles['Normal']

# Save
out_path = 'reports/tai-lieu-du-an-nhom-01.docx'
doc.save(out_path)
print('updated', out_path)
