from docx import Document
import os

path = os.path.join('docs', '2. Mau tai lieu.docx')
doc = Document(path)

# Replace placeholder Chapter 2 content (paragraphs 105-109 in inspection)
paragraph_updates = {
    105: 'Chương 2. Xây dựng pipeline và tiền xử lý dữ liệu',
    106: 'Chương này trình bày quá trình xây dựng pipeline dữ liệu và các bước tiền xử lý để chuyển dữ liệu thô từ nhiều bảng liên quan thành một tập dữ liệu sạch, nhất quán và sẵn sàng cho phân tích, feature engineering và huấn luyện mô hình.',
    107: '2.1. Quy trình tích hợp dữ liệu: Dữ liệu được thu thập từ nhiều bảng khác nhau, trong đó bảng application_train là bảng trung tâm và các bảng phụ như bureau, previous_application, installments_payments, POS_CASH_balance và credit_card_balance cung cấp các thông tin lịch sử tín dụng của khách hàng. Nhóm thiết kế pipeline theo hướng phân tầng bằng cách tổng hợp các bảng phụ thành các bảng summary ở mức khách hàng trước khi nối vào bảng chính bằng LEFT JOIN để tạo thành bảng application_flat. Cách tiếp cận này giúp giảm bộ nhớ, tối ưu hiệu suất xử lý và giữ được tính nhất quán của dữ liệu trong quá trình xử lý hàng triệu dòng.',
    108: '2.2. Làm sạch dữ liệu chuyên sâu: Sau khi dữ liệu đã được tích hợp, nhóm xử lý các vấn đề về missing value, sai logic và outlier. Một số giá trị NULL trong các nhóm cột summary phản ánh thực tế khách hàng chưa có lịch sử tương ứng, nên không được điền một cách máy móc. Ngược lại, với các cột đếm và flag, nhóm áp dụng quy tắc điền 0 và bổ sung cờ để giữ tín hiệu dịch vụ. Với các giá trị sai logic như DAYS_EMPLOYED = 365243 hoặc CODE_GENDER = XNA, nhóm phân tích ngữ cảnh nghiệp vụ và điều chỉnh về giá trị hợp lệ. Các biến tiền tệ và chỉ số lịch sử có phân phối lệch được xử lý bằng các biến đổi phù hợp như cắt theo phân vị hoặc log-transform.',
    109: '2.3. Kiểm thử dữ liệu và kiểm soát chất lượng: Pipeline được tích hợp các bước kiểm chứng để đảm bảo dữ liệu sau bước join, aggregate và clean không bị sai lệch. Nhóm kiểm tra số dòng, số cột, tính duy nhất của khóa khách hàng, kiểu dữ liệu và sự loại bỏ các giá trị sai logic. Các kịch bản kiểm thử này giúp phát hiện sớm lỗi dữ liệu và đảm bảo dữ liệu đầu vào có chất lượng cao trước khi đi vào phases EDA, feature engineering và huấn luyện mô hình.',
}

for idx, new_text in paragraph_updates.items():
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.text = new_text
        if idx == 105:
            p.style = doc.styles['Heading 1']
        else:
            p.style = doc.styles['Normal']
    else:
        print(f'Index {idx} out of range, doc has {len(doc.paragraphs)} paragraphs')

# Save changes
out_path = os.path.join('docs', '2. Mau tai lieu.docx')
doc.save(out_path)
print('Updated file:', out_path)
